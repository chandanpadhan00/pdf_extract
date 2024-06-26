import os
import re
import pdfplumber
from pdf2docx import Converter
from docx import Document
from docx2pdf import convert
from docx.shared import Pt

def pdf_to_docx(pdf_path, docx_path):
    cv = Converter(pdf_path)
    cv.convert(docx_path, start=0, end=None)
    cv.close()

def clean_docx(docx_path, cleaned_docx_path):
    doc = Document(docx_path)
    section_16_pattern = re.compile(r"16\s+HOW\s+SUPPLIED/STORAGE\s+AND\s+HANDLING", re.IGNORECASE)
    section_17_pattern = re.compile(r"17\s+PATIENT\s+COUNSELING\s+INFORMATION", re.IGNORECASE)
    
    new_doc = Document()
    copy = False

    for element in doc.element.body:
        if element.tag.endswith('p'):
            para = Document().paragraphs[0]._element
            para = element
            para = Document(new_doc._element, new_doc).paragraphs[-1]
            text = para.text

            if section_16_pattern.search(text):
                copy = True
            if section_17_pattern.search(text):
                copy = False
            if copy:
                new_paragraph = new_doc.add_paragraph(text)
                for run in para.runs:
                    new_run = new_paragraph.add_run(run.text)
                    new_run.bold = run.bold
                    new_run.italic = run.italic
                    new_run.underline = run.underline
                    new_run.font.size = run.font.size
                    new_run.font.name = run.font.name
        elif element.tag.endswith('tbl'):
            if copy:
                tbl = Document().tables[0]._element
                tbl = element
                new_tbl = new_doc.add_table(rows=0, cols=len(tbl.tr_lst[0].tc_lst))
                new_tbl.style = 'Table Grid'
                for row in tbl.tr_lst:
                    new_row = new_tbl.add_row().cells
                    for idx, cell in enumerate(row.tc_lst):
                        new_row[idx].text = cell.text

    new_doc.save(cleaned_docx_path)

def docx_to_pdf(docx_path, pdf_path):
    convert(docx_path, pdf_path)

def process_all_pdfs(input_folder, output_folder, start_reading_from=2):
    if not os.path.exists(output_folder):
        os.makedirs(output_folder)

    for filename in os.listdir(input_folder):
        if filename.endswith(".pdf"):
            pdf_path = os.path.join(input_folder, filename)
            base_filename = os.path.splitext(filename)[0]
            docx_path = os.path.join(output_folder, f"{base_filename}.docx")
            cleaned_docx_path = os.path.join(output_folder, f"{base_filename}_cleaned.docx")
            cleaned_pdf_path = os.path.join(output_folder, f"{base_filename}_16thSec.pdf")
            
            # Convert PDF to DOCX
            pdf_to_docx(pdf_path, docx_path)
            
            # Clean the DOCX
            clean_docx(docx_path, cleaned_docx_path)
            
            # Convert cleaned DOCX back to PDF
            docx_to_pdf(cleaned_docx_path, cleaned_pdf_path)
            
            print(f"Processed and cleaned {filename}")

# Main processing
input_folder = "path_to_your_input_folder"  # Update this path with the actual input folder path
output_folder = "path_to_your_output_folder"  # Update this path with the actual output folder path

process_all_pdfs(input_folder, output_folder, start_reading_from=2)
